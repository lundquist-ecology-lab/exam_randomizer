#%%
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Define the number of questions and choices
num_questions = 50
choices = ['a', 'b', 'c', 'd', 'e']

# Define the page size and margins
page_width, page_height = docx.shared.Inches(8.5), docx.shared.Inches(11)
left_margin, right_margin, top_margin, bottom_margin = docx.shared.Inches(0.5), docx.shared.Inches(0.5), docx.shared.Inches(0.5), docx.shared.Inches(0.5)
column_spacing, row_spacing = docx.shared.Inches(0.25), docx.shared.Inches(0.25)
bubble_size = docx.shared.Pt(12)

# Define the function for drawing bubbles
def draw_bubble(paragraph, is_selected):
    run = paragraph.add_run()
    run.font.size = bubble_size
    run.font.name = 'Arial'
    run.text = 'O' if is_selected else 'o'

# Create a new Word document and set the margins
doc = docx.Document()
sections = doc.sections
for section in sections:
    section.left_margin = left_margin
    section.right_margin = right_margin
    section.top_margin = top_margin
    section.bottom_margin = bottom_margin

# Loop over each question and draw the question number and answer choices
for i in range(num_questions):
    # Determine the row and column of the current question
    row = i // 6
    column = i % 6
    
    # Determine the x and y coordinates of the top-left corner of the question block
    x = left_margin + column * (page_width - left_margin - right_margin - 5 * bubble_size) // 6 + column_spacing * column
    y = top_margin + row * (page_height - top_margin - bottom_margin - 2 * bubble_size) // 8 + row_spacing * row
    
    # Add the question number
    question_number = str(i + 1) + '.'
    question_paragraph = doc.add_paragraph()
    question_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    question_paragraph.add_run(question_number).font.bold = True
    question_paragraph.add_run().text = '\t'
    
    # Add the answer choices
    answer_paragraph = doc.add_paragraph()
    answer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    answer_paragraph.add_run().text = '\t\t'
    for j, choice in enumerate(choices):
        draw_bubble(answer_paragraph, False)
        answer_paragraph.add_run('\t' + choice + '\t')
    answer_paragraph.add_run().text = '\t\t'
    
    # Add spacing between each question
    if i != num_questions - 1:
        doc.add_paragraph()

# Save the Word document
doc.save('bubble_sheet.docx')
# %%

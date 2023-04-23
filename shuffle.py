# %%
import re
import random

# Open the text file and read the lines
with open('yourfile.txt', 'r') as file:
    lines = file.readlines()

# Find all the questions and answers
questions = []
current_question = None
for line in lines:
    line = line.strip()
    if line.startswith('Q.'):
        current_question = line
        questions.append((current_question, []))
    elif line.startswith(('a.', 'b.', 'c.', 'd.', 'e.')):
        answer = line
        questions[-1][1].append(answer)

# Shuffle the questions and answers
random.shuffle(questions)
for i, question in enumerate(questions):
    answers = question[1]
    random.shuffle(answers)
    questions[i] = (question[0], *answers)

# Add the shuffled questions and answers to a new list
shuffled_questions = []
for i, question in enumerate(questions):
    shuffled_question = 'Q.' + str(i+1) + re.sub(r'^Q\.[\d.]+\s', ' ', question[0]) + '\n'
    for j, answer in enumerate(question[1:]):
        shuffled_question += chr(97+j) + '. ' + answer[3:] + '\n'
    shuffled_question += '\n'
    shuffled_questions.append(shuffled_question)

# Save the shuffled questions and answers to a new file
with open('shuffled.txt', 'w') as file:
    file.writelines(shuffled_questions)

# %%
import docx

# Read in the shuffled questions and answers from a text file
with open('shuffled.txt', 'r') as file:
    shuffled_questions = file.readlines()

# Create a new Word document
doc = docx.Document()

# Loop through the shuffled questions and answers and add them to the Word document
for question in shuffled_questions:
    # Remove any leading or trailing whitespace
    question = question.strip()
    # If the line starts with "Q.", it's a new question, so add a new paragraph for it
    if question.startswith('Q.'):
        doc.add_paragraph(question)
    # Otherwise, it's an answer, so add it to the current question's paragraph
    else:
        # Get the last paragraph in the document (which should be the current question's paragraph)
        current_question_para = doc.paragraphs[-1]
        # Add the answer as a new run (i.e., a new chunk of text within the same paragraph)
        current_question_para.add_run('\n' + question)

# Save the document as a new file
doc.save('shuffled_questions.docx')

# %%
import os

# specify the path of the file to be deleted
file_path = 'shuffled.txt'

# check if the file exists
if os.path.exists(file_path):
    # delete the file
    os.remove(file_path)
    print(f"The file {file_path} has been deleted.")
else:
    print(f"The file {file_path} does not exist.")
#%%